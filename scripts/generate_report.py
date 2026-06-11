"""Generate AGD Research Report as a Word document."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ──
for sec in doc.sections:
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(2.5)

# ── Helper functions ──
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = RGBColor(0x5B, 0x9B, 0xD5)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)
    return p

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.style.font.size = Pt(11)
    return p

def code_block(text):
    for line in text.strip().split('\n'):
        p = doc.add_paragraph()
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F2F2F2')
        p._p.get_or_add_pPr().append(shd)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        # Blue background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '2E74B5')
        tcPr.append(shd)
        # White bold text
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            table.rows[r_idx + 1].cells[c_idx].text = val
    doc.add_paragraph()

# ══════════════════════════════════════════════════════
#  TITLE PAGE
# ══════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_p.add_run('Adversarial Geometric Distillation (AGD) Pipeline')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub_p.add_run('EXP_04 — Complete Project Walkthrough & Research Proposal Coverage Report')
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('University of Moratuwa  |  Research in Generative AI  |  April 2026')

doc.add_page_break()

# ══════════════════════════════════════════════════════
#  SECTION 1 — WHAT IS THIS PROJECT?
# ══════════════════════════════════════════════════════
h1('1. What Is This Project About?')
body(
    'When AI generates 3D models from text prompts (e.g., "a knight in armor"), '
    'the results often have geometric defects called 3D hallucinations. Common examples include:'
)
bullet('Janus faces — a face appears on both the front AND back of the head')
bullet('Duplicated limbs — extra arms or legs growing from the body')
bullet('Floating blobs — random geometry disconnected from the main shape')
bullet('Holes and spikes — the surface tears or has thin protrusions')
body(
    'Your research proposes an Adversarial Geometric Distillation (AGD) system that '
    '(1) detects these hallucinations automatically using AI vision and mathematics, '
    '(2) locates exactly WHERE on the 3D model they occur, and '
    '(3) fixes them by nudging the geometry back into shape. '
    'The EXP_04 folder contains a working prototype of this system.'
)

# ══════════════════════════════════════════════════════
#  SECTION 2 — PROJECT STRUCTURE
# ══════════════════════════════════════════════════════
h1('2. Project File Structure')
add_table(
    ['File / Folder', 'Role'],
    [
        ['agd_pipeline.py', 'Main entry point — orchestrates the full pipeline'],
        ['discriminator.py', 'Topological + spectral analysis; hallucination score'],
        ['critic.py', 'Per-vertex anomaly critic (heuristic + GNN)'],
        ['gnn_train.py', 'GNN training script using pseudo-labels'],
        ['renderer.py', 'Multi-view PNG renderer (6 camera angles)'],
        ['view_consistency.py', 'Cross-view cosine similarity checker'],
        ['lmm_detector.py', 'LLaVA vision-language hallucination detector'],
        ['grounding.py', 'Maps 2D view signals → 3D vertex weights'],
        ['optimizer.py', 'Laplacian-based geometry refinement'],
        ['analysis_notebook.ipynb', 'Interactive Jupyter demo notebook'],
        ['requirements.txt', 'Python dependencies'],
        ['3d_samples/', '21 test meshes (bunny, dragon, janus, etc.)'],
        ['agd_outputs/', 'Refined meshes and rendered views'],
        ['LLaVA/', 'Locally cloned LLaVA multi-modal model'],
    ]
)

# ══════════════════════════════════════════════════════
#  SECTION 3 — PIPELINE STEPS
# ══════════════════════════════════════════════════════
h1('3. How the Pipeline Works (Step by Step)')
steps = [
    ('Step 1: Load Mesh', 'Load a .obj, .ply, or .stl file from the 3d_samples/ directory.'),
    ('Step 2: Baseline Analysis', 'Compute topology (Euler characteristic, genus, components) and spectral metrics (Fiedler value, eigenvalue variance).'),
    ('Step 3: Critic Scoring', 'Assign a per-vertex anomaly score using vertex degree and edge-length z-scores.'),
    ('Step 4: Multi-View Render', 'Render 6 views (front, back, left, right, top, bottom) as PNG images.'),
    ('Step 5: View Consistency', 'Compare opposite view pairs using cosine similarity to detect cross-view inconsistencies.'),
    ('Step 6: LLaVA Inspection (optional)', 'Run local LLaVA model on each view; parse JSON severity score.'),
    ('Step 7: Grounding', 'Map view-level severity scores to 3D vertex weights using projection + normal filtering.'),
    ('Step 8: Optimization', 'Apply weighted Laplacian smoothing for N iterations, guided by vertex weights.'),
    ('Step 9: Save & Report', 'Export refined mesh and print before/after hallucination scores.'),
]
for title, desc in steps:
    p = doc.add_paragraph()
    run = p.add_run(title + ':  ')
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p.add_run(desc)

body('')
body('CLI Usage:')
code_block(
    'python agd_pipeline.py --input-dir 3d_samples --output-dir agd_outputs --render-views\n'
    '\n'
    '# With LLaVA vision model:\n'
    'python agd_pipeline.py --input-dir 3d_samples --output-dir agd_outputs \\\n'
    '  --render-views --use-vlm --llava-model-path liuhaotian/llava-v1.5-7b'
)

# ══════════════════════════════════════════════════════
#  SECTION 4 — FILE-BY-FILE
# ══════════════════════════════════════════════════════
h1('4. File-by-File Explanation with Code')

# 4.1 agd_pipeline.py
h2('4.1  agd_pipeline.py — The Conductor (164 lines)')
body(
    'Ties every module together. Scans a folder for mesh files and runs the full '
    'detect → ground → refine → evaluate loop for each mesh.'
)
code_block(
    'mesh = trimesh.load(path, force="mesh")\n'
    'before_topo, before_spec, before_score = analyse(mesh)\n'
    'scores = critic.score(mesh)\n'
    'view_paths = render_views(mesh, view_dir)\n'
    'consistency = check_view_consistency(view_paths, similarity_threshold=0.85)\n'
    'weights = build_region_weights(mesh, scores, bias_weights=bias)\n'
    'refined = refine_mesh(mesh, weights=weights, iterations=10, lr=0.15)\n'
    'after_topo, after_spec, after_score = analyse(refined)\n'
    'print(f"score: {before_score:.3f} -> {after_score:.3f}")'
)

# 4.2 discriminator.py
h2('4.2  discriminator.py — The Math Inspector (358 lines)')
body(
    'Analyses the mesh using two mathematical lenses:'
)
h3('A) Topological Analysis')
body('Uses the Euler characteristic formula:  χ = V − E + F   and  genus g = (2 − χ) / 2')
code_block(
    'def compute_topology(mesh):\n'
    '    components = mesh.split(only_watertight=False)\n'
    '    euler_char = num_vertices - num_edges + num_faces\n'
    '    genus = (2.0 - euler_char) / 2.0\n'
    '    return {"euler_char": euler_char, "genus": genus,\n'
    '            "components": len(components)}'
)
h3('B) Spectral Analysis')
body('Builds the graph Laplacian L = D − A and analyses eigenvalues for irregularity.')
code_block(
    'L = nx.laplacian_matrix(subgraph).astype(float)\n'
    'trace_l  = float(L.diagonal().sum())\n'
    'trace_l2 = float(L.multiply(L).sum())\n'
    'var = trace_l2 / n - (trace_l / n) ** 2   # eigenvalue variance\n'
    'eigenvalues = np.linalg.eigvalsh(L.toarray())\n'
    'fiedler = eigenvalues[1]                   # algebraic connectivity'
)
h3('C) Hallucination Score')
body('Combines topology and spectral signals into a 0–1 score.')
code_block(
    'score = (0.25 * comp_penalty    # multiple disconnected pieces\n'
    '       + 0.35 * genus_penalty   # surface holes\n'
    '       + 0.25 * var_penalty     # spectral irregularity\n'
    '       + 0.15 * fiedler_penalty)# weak connectivity'
)

# 4.3 critic.py
h2('4.3  critic.py — Per-Vertex Anomaly Judge (140 lines)')
body('Assigns an anomaly score to EVERY vertex. Two implementations:')
h3('Heuristic Critic (works immediately)')
code_block(
    'deg_z = zscore(mesh.vertex_degree)          # unusual connections?\n'
    'len_z = zscore(mean_edge_lengths)           # unusual edge lengths?\n'
    'score = 0.5*sigmoid(|deg_z|) + 0.5*sigmoid(|len_z|)'
)
h3('GNN Critic (requires PyTorch Geometric)')
code_block(
    'class _SimpleGCN:\n'
    '    def __init__(self):\n'
    '        self.conv1 = GCNConv(2, 16)  # input: [degree, avg_edge_len]\n'
    '        self.conv2 = GCNConv(16, 1)  # output: anomaly score per vertex'
)

# 4.4 gnn_train.py
h2('4.4  gnn_train.py — Teaching the GNN (119 lines)')
body('Trains the GNN using pseudo-labels generated by the heuristic critic.')
code_block(
    'scores = heuristic_critic.score(mesh)\n'
    'labels = (scores >= 0.7).astype(float)  # 1 = hallucinated\n'
    '\n'
    'for epoch in range(args.epochs):\n'
    '    out  = model(batch.x, batch.edge_index)\n'
    '    loss = BCEWithLogitsLoss(out, batch.y)\n'
    '    loss.backward(); optimizer.step()\n'
    '\n'
    'torch.save(model.state_dict(), "gnn_critic.pt")'
)
body('Run with:')
code_block('python gnn_train.py --input-dir 3d_samples --output gnn_critic.pt --epochs 20')

# 4.5 renderer.py
h2('4.5  renderer.py — Multi-View Camera (64 lines)')
body('Renders 6 views of a mesh as PNG images using trimesh scene rendering.')
code_block(
    'DEFAULT_VIEWS = {\n'
    '    "front":  (0, 0, 0),      "back":   (0, π, 0),\n'
    '    "left":   (0, π/2, 0),    "right":  (0, -π/2, 0),\n'
    '    "top":    (-π/2, 0, 0),   "bottom": (π/2, 0, 0),\n'
    '}\n'
    'for name, angles in DEFAULT_VIEWS.items():\n'
    '    scene.set_camera(angles=angles, distance=extent*2.5)\n'
    '    png = scene.save_image(resolution=(512, 512))'
)

# 4.6 view_consistency.py
h2('4.6  view_consistency.py — Symmetry Checker (65 lines)')
body(
    'Compares opposite view pairs (front↔back, left↔right, top↔bottom) '
    'using cosine similarity on 32×32 grayscale embeddings.'
)
code_block(
    'def _embed_image(path, size=32):\n'
    '    img = Image.open(path).convert("L").resize((size, size))\n'
    '    arr = np.asarray(img).reshape(-1)\n'
    '    return arr / np.linalg.norm(arr)\n'
    '\n'
    'sim = cosine(embed("front"), embed("back"))\n'
    'if sim < 0.85:\n'
    '    severity = 1.0 - sim   # low similarity = higher severity'
)

# 4.7 lmm_detector.py
h2('4.7  lmm_detector.py — AI Vision Inspector (174 lines)')
body('Uses a local LLaVA model to visually inspect each rendered view and return a severity score.')
code_block(
    'query = (\n'
    '    "You are checking a 3D render for geometric hallucinations. "\n'
    '    f"The view name is \'{view_name}\'. "\n'
    '    "Return JSON: {severity: 0-1, notes: string}"\n'
    ')\n'
    '# Model returns e.g.: {"severity": 0.3, "notes": "duplicated face"}\n'
    'severity, notes = _parse_detection(model_output)'
)

# 4.8 grounding.py
h2('4.8  grounding.py — 3D GPS (93 lines)')
body(
    'KEY RESEARCH CONTRIBUTION: Converts 2D view-level detections into '
    '3D vertex-level weights. This is the "semantic bottleneck removal" step.'
)
code_block(
    'for view, score in view_scores.items():\n'
    '    direction = directions[view]  # e.g., "back" -> [0, 0, -1]\n'
    '    projection = (vertices - center) @ direction\n'
    '    threshold  = quantile(projection, 1 - view_extent)  # top 30%\n'
    '    facing     = (normals @ direction) >= 0.2\n'
    '    mask       = (projection >= threshold) & facing\n'
    '    bias[mask] = max(bias[mask], score)\n'
    '\n'
    'weights = build_region_weights(mesh, critic_scores,\n'
    '    top_fraction=0.1, expand_hops=1, bias_weights=bias)'
)

# 4.9 optimizer.py
h2('4.9  optimizer.py — Geometry Fixer (44 lines)')
body('Moves vertices to reduce hallucinations using two competing forces:')
bullet('Laplacian smoothing (L_geo) — pull vertex toward average of neighbors')
bullet('Anchor term (L_distill) — keep vertex close to original position')
code_block(
    'for _ in range(iterations):\n'
    '    lap     = uniform_laplacian(vertices, neighbors)  # v_i - mean(N(i))\n'
    '    distill = vertices - anchor\n'
    '    step    = lambda_geo * lap + lambda_distill * distill\n'
    '    vertices = vertices - lr * step * weights  # only move "bad" vertices!'
)
body('Formula:  v_i ← v_i − η·(λ·∇L_geo + (1−λ)·∇L_distill) ⊙ w_i')
body(
    'where w_i ≈ 1.0 for hallucinated vertices and w_i ≈ 0.1 for clean vertices.'
)

# ══════════════════════════════════════════════════════
#  SECTION 5 — DATASET
# ══════════════════════════════════════════════════════
h1('5. Test Dataset (3d_samples/ — 21 Meshes)')
add_table(
    ['File', 'Format', 'Size', 'Purpose'],
    [
        ['bunny.ply', 'PLY', '2.9 MB', 'Stanford Bunny — clean reference mesh'],
        ['dragon.ply', 'PLY', '32 MB', 'Stanford Dragon — complex clean mesh'],
        ['janus.obj', 'OBJ', '1.9 MB', 'Known Janus-face artifact for testing'],
        ['horse.ply', 'PLY', '2.1 MB', 'Animal mesh'],
        ['happy.ply', 'PLY', '40 MB', 'Stanford Happy Buddha'],
        ['hand.ply', 'PLY', '31 MB', 'Detailed hand mesh'],
        ['blade.ply', 'PLY', '80 MB', 'Complex weapon mesh'],
        ['Mr_Bean.obj', 'OBJ', '186 KB', 'Character face mesh'],
        ['elepham.obj', 'OBJ', '2.8 MB', 'Animal mesh'],
        ['venusm.obj', 'OBJ', '3.1 MB', 'Classical sculpture'],
        ['ateneam.obj', 'OBJ', '1 MB', 'Classical sculpture'],
        ['Meshy_AI_Ironclad.obj', 'OBJ', '76 MB', 'AI-generated mesh (Meshy AI)'],
        ['Bearded_guy.ply', 'PLY', '37 MB', 'Character mesh'],
        ['Rigged_Hand.obj', 'OBJ', '144 KB', 'Rigged hand model'],
        ['mba1.obj / mba2.obj', 'OBJ', '13/11 MB', 'Additional test meshes'],
        ['*.tri files (5)', 'TRI', 'Various', 'Triangle mesh format samples'],
    ]
)

# ══════════════════════════════════════════════════════
#  SECTION 6 — COVERAGE ASSESSMENT
# ══════════════════════════════════════════════════════
h1('6. Research Proposal Coverage Assessment')
body('Overall estimated coverage: ~60–65% of the full research proposal.')
doc.add_paragraph()

add_table(
    ['#', 'Proposal Component', 'Status', 'Notes'],
    [
        ['1', 'SDS-based text-to-3D generation', '❌ Not implemented', 'Pipeline starts from existing meshes'],
        ['2', 'Multi-view rendering', '✅ Implemented', '6 standard views via trimesh (renderer.py)'],
        ['3', 'LMM hallucination detection', '✅ Implemented', 'Local LLaVA with JSON severity output'],
        ['4', 'View consistency checking', '✅ Implemented', 'Cosine similarity embeddings'],
        ['5', 'Adversarial geometric critic', '🟡 Partial', 'Heuristic works; GNN is a scaffold'],
        ['6', 'Topological metrics (Euler, genus, Betti)', '✅ Implemented', 'Full analysis in discriminator.py'],
        ['7', 'Spectral analysis (Laplacian)', '✅ Implemented', 'Fiedler value, variance'],
        ['8', 'Geometric grounding (2D → 3D)', '✅ Implemented', 'Projection + normal filtering'],
        ['9', 'Energy-based refinement loop', '✅ Implemented', 'Laplacian + anchor in optimizer.py'],
        ['10', 'Closed-loop pipeline', '✅ Implemented', 'Full orchestration in agd_pipeline.py'],
        ['11', 'Test dataset', '✅ Provided', '21 meshes including janus.obj artifact'],
        ['12', 'VLM spatial localization (bbox/mask)', '❌ Not implemented', 'LLaVA gives severity only'],
        ['13', 'Differentiable topology losses', '❌ Not implemented', 'Metrics computed but not as loss'],
        ['14', 'Benchmark evaluation (Objaverse/GSO)', '❌ Not implemented', 'No benchmark scripts'],
        ['15', 'Comparison with baselines', '❌ Not implemented', 'No comparison framework'],
    ]
)

h2('Coverage by Proposal Block')
add_table(
    ['Block', 'Description', 'Coverage', 'Confidence'],
    [
        ['Block 1: Detection', 'LMM + multi-view + topology', '~80%', 'High — two detectors work'],
        ['Block 2: Critic', 'GNN + adversarial training', '~40%', 'Low — GNN needs real supervision'],
        ['Block 3: Grounding', 'Map 2D signals → 3D constraints', '~70%', 'Medium — coarse mapping works'],
        ['Block 4: Refinement', 'Optimize geometry with energy', '~75%', 'Medium — mesh-only, no SDS'],
        ['Block 5: Evaluation', 'Benchmarks + comparisons', '~20%', 'Low — metrics exist, no benchmarks'],
    ]
)

# ══════════════════════════════════════════════════════
#  SECTION 7 — STRENGTHS & GAPS
# ══════════════════════════════════════════════════════
h1('7. Strengths and Gaps')

h2('What Is Working Well')
bullet('Complete loop structure — detect → ground → refine pipeline runs end-to-end')
bullet('Modular design — each component is a separate file with clean interfaces')
bullet('Dual detection — both mathematical (discriminator) and visual (LLaVA) detection')
bullet('Grounding innovation — 2D-to-3D weight mapping is the core research contribution')
bullet('Rich test dataset — 21 meshes including known artifacts for validation')
bullet('Scalable analysis — spectral analysis auto-samples large meshes (>2000 vertices)')

h2('What Needs Work')
bullet('No text-to-3D generation — must integrate with SDS/Threestudio/NeRF pipeline')
bullet('GNN critic is toy-level — needs real annotated data or adversarial training')
bullet('LLaVA gives only severity — no spatial localization (bounding boxes or masks)')
bullet('Topology not used as loss — Betti numbers computed but not optimized against')
bullet('No quantitative benchmarks — need Objaverse/GSO evaluation + comparison tables')

# ══════════════════════════════════════════════════════
#  SECTION 8 — NEXT STEPS
# ══════════════════════════════════════════════════════
h1('8. Recommended Next Steps (Priority Order)')
steps_next = [
    ('1. Connect SDS generation',
     'Integrate Threestudio so the full text → 3D → refine loop works.'),
    ('2. Make topology differentiable',
     'Use persistent homology libraries (e.g., giotto-tda) as a differentiable loss term.'),
    ('3. Upgrade LLaVA grounding',
     'Use vision tower features or grounding models for precise 2D→3D localization.'),
    ('4. Adversarial critic training',
     'Train GNN critic to distinguish clean vs. hallucinated, then train refinement to fool it.'),
    ('5. Benchmark scripts',
     'Write evaluation on Objaverse subset with automated metric tables for paper comparison.'),
]
for title, desc in steps_next:
    p = doc.add_paragraph()
    run = p.add_run(title + ':  ')
    run.bold = True
    p.add_run(desc)

# ══════════════════════════════════════════════════════
#  SECTION 9 — DEPENDENCIES
# ══════════════════════════════════════════════════════
h1('9. Dependencies (requirements.txt)')
add_table(
    ['Package', 'Version', 'Used For'],
    [
        ['numpy', '>=1.20', 'Array math throughout'],
        ['trimesh', '>=4.10.0', 'Mesh loading, topology, rendering'],
        ['networkx', '>=2.5', 'Graph Laplacian construction'],
        ['scipy', '>=1.10', 'Sparse eigenvalue solver'],
        ['pillow', '>=10.0', 'Image I/O for view rendering'],
        ['torch', '>=2.2', 'GNN critic and training'],
        ['torch-geometric', '>=2.5', 'GCNConv layers for GNN critic'],
        ['transformers', '>=4.38', 'LLaVA model loading'],
        ['accelerate', '>=0.27', 'LLaVA inference acceleration'],
        ['pyglet', '<2', 'Trimesh offline rendering (must be v1.x)'],
    ]
)
body('Install with:')
code_block('pip install -r requirements.txt\npip install -e LLaVA')

# ══════════════════════════════════════════════════════
#  SAVE
# ══════════════════════════════════════════════════════
out_path = os.path.join(
    r'c:\Users\ASUS\OneDrive - University of Moratuwa\Desktop\Research in Gen AI\EXP_04',
    'AGD_Research_Report.docx'
)
doc.save(out_path)
print(f"Saved: {out_path}")
