from __future__ import annotations

import argparse
import shutil
import tempfile
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from lxml import etree

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W_NS}


def set_style_font(style, *, size: int, bold: bool, before: int = 0, after: int = 0) -> None:
    style.font.name = "Times New Roman"
    style.font.size = Pt(size)
    style.font.bold = bold
    pf = style.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.keep_with_next = bold


def clear_run_format(paragraph) -> None:
    for run in paragraph.runs:
        if not run.text:
            continue
        run.style = None
        run.font.name = None
        run.font.size = None
        run.font.bold = None


def insert_paragraph_after(paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def add_toc_field(paragraph, levels: str = "1-3") -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f' TOC \\o "{levels}" \\h \\z '

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "(TOC will populate after updating fields)"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(placeholder)
    run._r.append(fld_end)


def set_update_fields_on_open(docx_path: Path) -> None:
    tmp = Path(tempfile.mkdtemp(prefix="docx_update_fields_"))
    try:
        with zipfile.ZipFile(docx_path, "r") as zf:
            zf.extractall(tmp)

        settings_path = tmp / "word" / "settings.xml"
        settings_path.parent.mkdir(parents=True, exist_ok=True)
        parser = etree.XMLParser(remove_blank_text=False)

        if settings_path.exists():
            tree = etree.parse(str(settings_path), parser)
            root = tree.getroot()
        else:
            root = etree.Element(f"{{{W_NS}}}settings", nsmap={"w": W_NS})
            tree = etree.ElementTree(root)

        update_fields = root.find("w:updateFields", namespaces=NS)
        if update_fields is None:
            update_fields = etree.Element(f"{{{W_NS}}}updateFields")
            root.insert(0, update_fields)
        update_fields.set(f"{{{W_NS}}}val", "true")

        tree.write(str(settings_path), xml_declaration=True, encoding="UTF-8", standalone="yes")

        with zipfile.ZipFile(docx_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            for file_path in tmp.rglob("*"):
                if file_path.is_dir():
                    continue
                zf.write(file_path, file_path.relative_to(tmp).as_posix())
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def normalize_styles(doc: Document) -> None:
    set_style_font(doc.styles["Heading 1"], size=14, bold=True, before=12, after=6)
    set_style_font(doc.styles["Heading 2"], size=12, bold=True, before=10, after=4)
    set_style_font(doc.styles["Heading 3"], size=12, bold=True, before=8, after=3)

    if "TOC Heading" in [style.name for style in doc.styles]:
        set_style_font(doc.styles["TOC Heading"], size=14, bold=True, before=12, after=6)

    for style_name in ["Body Text", "First Paragraph", "List Paragraph", "Compact", "Bibliography"]:
        if style_name in [style.name for style in doc.styles]:
            set_style_font(doc.styles[style_name], size=12, bold=False, before=0, after=0)

    for style_name in ["TOC 1", "TOC 2", "TOC 3"]:
        if style_name in [style.name for style in doc.styles]:
            set_style_font(doc.styles[style_name], size=12, bold=False, before=0, after=0)


def apply_heading_map(doc: Document) -> None:
    chapter_1_h2 = {
        "Introduction",
        "Objectives",
        "Problem in Brief",
        "Background Motivation",
        "Noval Approach to Mitigating Hallucinations in 3D Model Generation",
        "Resource Requirements",
        "Structure of the Thesis",
        "Summary",
    }
    chapter_1_h3 = {"Hardware Requirements", "Software Requirements"}

    chapter_2_h3 = {
        "Early development of Generative 3D Assets",
        "Recent developments and Future trends",
        "Issues and research challenges",
        "Definition of the research gap/problem",
    }

    chapter_4_h2 = {
        "Introduction",
        "Hypothesis and its inspiration",
        "Inputs and Outputs of Extension",
        "Process workflow for Extension",
        "Technology Identified",
        "Features of technological extension",
        "Target Users / Use Scenarios",
        "Positioning the Extension within the AI body of knowledge",
        "Summary",
    }
    chapter_4_h3 = {
        "Stage 1 - Mesh Ingestion and Pre-processing",
        "Stage 2 - Dense Multi-View Rendering",
        "Stage 3 - LMM Hallucination Detection",
        "Stage 4 - Adversarial Geometric Critic",
        "Stage 5 - Geometric Grounding",
        "Stage 6 - Grounded Refinement Loop",
        "Target User Group 1 - 3D Generative AI Researchers",
        "Target User Group 2 - 3D Asset Creators (VR/AR/Gaming)",
        "Target User Group 3 - Medical and Engineering 3D Reconstruction",
        "Target User Group 4 - Academic Researchers in Geometric Deep Learning",
        "Relation to Generative AI (Text-to-3D)",
        "Relation to Multi-Modal AI",
        "Relation to Geometric Deep Learning",
        "Relation to Computational Geometry and Topology",
        "Novel Contribution",
    }

    current_chapter = None
    content_started = False
    style_names = {style.name for style in doc.styles}

    for paragraph in doc.paragraphs:
        text = (paragraph.text or "").strip()
        if not text:
            continue

        if text == "Declaration":
            content_started = True

        if text == "Table of Contents" and "TOC Heading" in style_names:
            paragraph.style = doc.styles["TOC Heading"]
        elif paragraph.style.name == "Heading 1" and text.startswith("Chapter "):
            current_chapter = text
        elif current_chapter == "Chapter 1" and text in chapter_1_h2:
            paragraph.style = doc.styles["Heading 2"]
        elif current_chapter == "Chapter 1" and text in chapter_1_h3:
            paragraph.style = doc.styles["Heading 3"]
        elif current_chapter == "Chapter 2" and text in chapter_2_h3:
            paragraph.style = doc.styles["Heading 3"]
        elif current_chapter == "Chapter 4" and text in chapter_4_h2:
            paragraph.style = doc.styles["Heading 2"]
        elif current_chapter == "Chapter 4" and text in chapter_4_h3:
            paragraph.style = doc.styles["Heading 3"]
        elif paragraph.style.name == "List Paragraph" and "Body Text" in style_names:
            paragraph.style = doc.styles["Body Text"]

        if (
            content_started
            and paragraph.style.name not in {"Heading 1", "Heading 2", "Heading 3", "TOC Heading", "Bibliography"}
            and "Body Text" in style_names
        ):
            paragraph.style = doc.styles["Body Text"]

        if not content_started:
            continue
        clear_run_format(paragraph)


def insert_toc(doc: Document) -> None:
    paragraphs = doc.paragraphs
    toc_heading_idx = None
    for idx, paragraph in enumerate(paragraphs):
        if (paragraph.text or "").strip() == "Table of Contents":
            toc_heading_idx = idx
            break

    if toc_heading_idx is None:
        return

    section_paragraphs = doc.paragraphs
    end_idx = None
    for idx in range(toc_heading_idx + 1, len(section_paragraphs)):
        paragraph = section_paragraphs[idx]
        if paragraph.style.name == "Heading 1" and (paragraph.text or "").strip():
            end_idx = idx
            break
    if end_idx is None:
        end_idx = len(section_paragraphs)

    for paragraph in list(section_paragraphs[toc_heading_idx + 1 : end_idx]):
        delete_paragraph(paragraph)

    toc_paragraph = insert_paragraph_after(doc.paragraphs[toc_heading_idx])
    add_toc_field(toc_paragraph, levels="1-3")


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input_docx", type=Path)
    parser.add_argument("output_docx", type=Path)
    args = parser.parse_args()

    if not args.input_docx.exists():
        raise FileNotFoundError(args.input_docx)

    doc = Document(str(args.input_docx))
    normalize_styles(doc)
    apply_heading_map(doc)
    args.output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.output_docx))
    set_update_fields_on_open(args.output_docx)
    print(args.output_docx)


if __name__ == "__main__":
    main()
