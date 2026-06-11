"""Generate the standalone Chapter 2 literature review DOCX for AGD."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT_NAME = "AGD_Literature_Review_Chapter_2.docx"


def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    r_fonts = run._element.rPr.rFonts
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), name)


def apply_font_to_paragraph(paragraph, name="Times New Roman", size=12, bold=False, italic=False):
    for run in paragraph.runs:
        set_run_font(run, name=name, size=size, bold=bold, italic=italic)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(4.0)
    section.right_margin = Cm(2.5)
    section.start_type = WD_SECTION.NEW_PAGE

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.5

    for style_name in ("Heading 1", "Heading 2"):
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.bold = True
        style.font.size = Pt(14 if style_name == "Heading 1" else 12)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True


def add_center_title(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_run_font(run, size=14, bold=True)
    return paragraph


def add_body(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, size=12)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 10)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(run, size=14 if level == 1 else 12, bold=True)
    return paragraph


def add_equation(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    set_run_font(run, size=11, italic=True)
    return paragraph


def set_table_header_repeat(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_alignment_center(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    jc = tbl_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        tbl_pr.append(jc)
    jc.set(qn("w:val"), "center")


def set_cell_text(cell, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_comparison_table(doc):
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(6)
    caption.paragraph_format.space_after = Pt(6)
    run = caption.add_run("Table 2.1 Comparison of representative methods for text-to-3D generation and hallucination mitigation")
    set_run_font(run, size=12, bold=True)

    headers = [
        "Study/Model",
        "Representation",
        "Core idea",
        "Strengths",
        "Limitations",
        "Relevance to AGD",
    ]
    rows = [
        [
            "Point Set Generation Network [7]",
            "Point cloud",
            "Predict unordered 3D points from a single image.",
            "Simple and geometric.",
            "No surface connectivity or semantics.",
            "Shows why later methods needed stronger structure.",
        ],
        [
            "NeRF [1]",
            "Implicit field",
            "Represent a scene as continuous color and density.",
            "Excellent novel-view synthesis.",
            "Needs multi-view supervision and is not generative by itself.",
            "Provides the field-and-rendering basis used by later text-to-3D systems.",
        ],
        [
            "Dream Fields [9]",
            "CLIP-guided NeRF",
            "Optimize a NeRF so renders match a text prompt.",
            "First zero-shot text-to-3D direction.",
            "Coarse geometry and weak multi-view consistency.",
            "Establishes the pre-SDS baseline for AGD.",
        ],
        [
            "DreamFusion [2]",
            "SDS over NeRF",
            "Use a frozen 2D diffusion prior to guide 3D optimization.",
            "Strong semantics without 3D training data.",
            "Janus artifacts and geometry drift.",
            "Primary baseline that motivates hallucination mitigation.",
        ],
        [
            "Latent-NeRF [12]",
            "Latent SDS",
            "Run text-to-3D optimization in the latent space of diffusion features.",
            "Better texture control than pixel-space optimization.",
            "Still inherits 2D-prior ambiguity.",
            "Illustrates that latent guidance alone does not solve structure.",
        ],
        [
            "Magic3D [3]",
            "Coarse-to-fine NeRF to mesh",
            "Use a low-resolution stage then refine a higher-quality mesh.",
            "Higher resolution and faster refinement.",
            "Improves fidelity more than topology.",
            "Relevant to AGD as a strong generation backbone that still needs correction.",
        ],
        [
            "Fantasia3D [4]",
            "Disentangled mesh optimization",
            "Separate geometry from appearance using normal-guided refinement.",
            "Better mesh extraction and material control.",
            "Geometry remains tied to 2D supervision.",
            "Shows the value of geometry-aware cues but not grounded error localization.",
        ],
        [
            "ProlificDreamer [5]",
            "Variational SDS",
            "Model 3D generation as a distribution rather than a single sample.",
            "Improves diversity and stability.",
            "High computation and no explicit topology reasoning.",
            "Important for AGD because diversity does not guarantee structural correctness.",
        ],
        [
            "Perp-Neg and ESD [16], [17]",
            "Score-space mitigation",
            "Modify guidance so side-view negatives or entropy regularization reduce mode collapse.",
            "Directly targets Janus symptoms.",
            "Geometry blind and indirect.",
            "Motivates AGD's shift from score shaping to spatial correction.",
        ],
        [
            "Hallo3D [6]",
            "LMM-in-the-loop correction",
            "Detect hallucinations with a multimodal model and issue negative prompts.",
            "Brings semantic reasoning into correction.",
            "Text feedback is coarse and may hallucinate.",
            "The main prior system that AGD seeks to extend with geometric grounding.",
        ],
    ]

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    set_table_alignment_center(table)

    widths = [Cm(2.3), Cm(2.0), Cm(2.5), Cm(2.2), Cm(2.6), Cm(2.7)]

    for column_index, header in enumerate(headers):
        cell = table.rows[0].cells[column_index]
        cell.width = widths[column_index]
        set_cell_text(cell, header, bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    set_table_header_repeat(table.rows[0])

    for row_index, row_data in enumerate(rows, start=1):
        row = table.rows[row_index]
        for col_index, value in enumerate(row_data):
            cell = row.cells[col_index]
            cell.width = widths[col_index]
            alignment = WD_ALIGN_PARAGRAPH.CENTER if col_index in (0, 1) else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cell, value, size=9, align=alignment)

    doc.add_paragraph()


def add_reference_entry(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.left_indent = Cm(0.64)
    paragraph.paragraph_format.first_line_indent = Cm(-0.64)
    run = paragraph.add_run(text)
    set_run_font(run, size=12)
    return paragraph


def build_document():
    doc = Document()
    configure_document(doc)

    add_center_title(doc, "CHAPTER 2")
    add_center_title(doc, "LITERATURE REVIEW")
    add_center_title(doc, "DEVELOPMENTS AND CHALLENGES IN TEXT-TO-3D GENERATION AND HALLUCINATION MITIGATION")
    doc.add_paragraph()

    add_heading(doc, "2.1 Introduction", level=1)
    add_body(
        doc,
        "The introductory chapter of this study argues that modern text-to-3D systems can produce visually appealing assets while still suffering from geometric hallucinations such as duplicated parts, floating components, broken symmetry, and weak manifold structure. That concern emerges from the historical trajectory of the field itself: neural scene representations such as NeRF [1] made continuous 3D reconstruction practical, while score-distilled pipelines such as DreamFusion [2], Magic3D [3], Fantasia3D [4], and ProlificDreamer [5] transferred the power of large 2D priors into 3D generation without giving those priors an explicit understanding of topology or object integrity."
    )
    add_body(
        doc,
        "The literature therefore needs to be read not only as a sequence of quality improvements, but also as a sequence of unresolved structural compromises. Recent systems such as Hallo3D [6] acknowledge this limitation by introducing multimodal hallucination detection, yet they still convert spatial errors into textual feedback and then ask a 2D prior to interpret that feedback again. For a problem like Adversarial Geometric Distillation (AGD), the most relevant question is not simply which model looks sharper, but which class of method can detect, localize, and correct structural defects in a way that is faithful to 3D geometry."
    )
    add_body(
        doc,
        "Accordingly, this chapter reviews prior work through an AGD-oriented lens. It first traces the chronological shift from explicit geometric representations to implicit fields and diffusion-guided generation, then compares major methods, critiques their strengths and weaknesses, and finally identifies the unresolved research gap that motivates a closed-loop, spatially grounded refinement framework."
    )

    add_heading(doc, "2.2 Chronological Review of Prior Work", level=1)
    add_heading(doc, "2.2.1 Early development of 3D generative representations", level=2)
    add_body(
        doc,
        "Early learning-based 3D generation relied on explicit representations such as point clouds and sparse voxels. Fan et al. [7] showed that a network could predict an unordered point set from a single image using geometric losses such as Chamfer Distance and Earth Mover's Distance, but the absence of explicit surface connectivity limited the method's ability to enforce watertight structure. Sparse voxel approaches later improved spatial occupancy modeling; for example, VoxGRAF [8] demonstrated that sparse voxel grids could support 3D-aware synthesis more efficiently than dense volumetric grids, although voxel methods still faced memory and resolution trade-offs."
    )
    add_body(
        doc,
        "A decisive shift occurred when continuous implicit neural fields replaced discrete geometry. NeRF [1] models a scene as a continuous radiance field that maps spatial location and viewing direction into color and density:"
    )
    add_equation(doc, "F_theta(x, d) -> (c, sigma)")
    add_body(
        doc,
        "Novel views are rendered by integrating this radiance along each camera ray:"
    )
    add_equation(doc, "C(r) = integral_(t_n)^(t_f) T(t) sigma(r(t)) c(r(t), d) dt")
    add_body(
        doc,
        "This formulation solved view synthesis elegantly, but it still assumed multi-view supervision and did not by itself provide an open-ended generative mechanism. Dream Fields [9] extended the idea by optimizing a NeRF so that rendered views matched a text prompt in CLIP space, thereby demonstrating zero-shot text-guided 3D generation. Even so, CLIP-guided supervision often produced coarse shapes, incomplete backsides, and weak consistency under novel viewpoints [9]."
    )
    add_body(
        doc,
        "Other early systems explored explicit generation routes with different trade-offs. Point-E [10] prioritized speed by producing point clouds directly from text prompts, while Shap-E [11] learned to generate the parameters of implicit functions and textured meshes more compactly than point-based pipelines. These methods broadened the design space, but they did not eliminate the central difficulty that AGD targets: how to preserve or recover global geometric coherence when generation is driven by incomplete 2D or text supervision."
    )

    add_heading(doc, "2.2.2 Recent developments and future trends", level=2)
    add_body(
        doc,
        "The current era of text-to-3D generation is dominated by diffusion-guided optimization. DreamFusion [2] introduced Score Distillation Sampling (SDS), which uses a frozen 2D diffusion model to supply a gradient signal for updating a 3D representation. In simplified form, the SDS gradient can be written as:"
    )
    add_equation(doc, "nabla_phi L_SDS = E_(t, eps)[ w(t) (eps_theta(x_t; y, t) - eps) * (partial x / partial phi) ]")
    add_body(
        doc,
        "This innovation avoided the need for large paired text-to-3D datasets, but it also inherited the biases of 2D priors. Subsequent methods refined the pipeline rather than replacing its basic logic. Latent-NeRF [12] moved the optimization into the latent space of the diffusion model, Magic3D [3] adopted a coarse-to-fine strategy that refined a high-quality mesh after an initial low-resolution stage, and Fantasia3D [4] separated geometry from appearance so that surface normals and materials could be controlled more deliberately."
    )
    add_body(
        doc,
        "More recent work has tried to improve either diversity or mesh quality. ProlificDreamer [5] reformulated generation as a distribution-level problem through Variational Score Distillation, helping the model escape some single-solution failure modes. MeshGPT [13] took a different route by treating triangle meshes as a sequence modeling problem, thereby moving toward mesh-native generation with cleaner topology. Parallel trends such as the resurgence of 3D Gaussian splatting [14] and symmetry-aware modeling in SYM3D [15] show that the field is actively searching for representations and inductive biases that reduce ambiguity, accelerate inference, and produce more controllable structure."
    )
    add_body(
        doc,
        "These trends point toward a future in which fidelity, efficiency, and structure are optimized together rather than traded against each other. However, the literature also shows that faster rendering, better textures, or stronger semantic alignment do not automatically resolve hallucinations. The field is still strongest at generating plausible appearances and weaker at guaranteeing that those appearances correspond to consistent 3D objects."
    )

    add_heading(doc, "2.2.3 Issues and research challenges", level=2)
    add_body(
        doc,
        "The most persistent challenge is that 2D priors are not intrinsically 3D-aware. DreamFusion-style pipelines can converge to high-probability frontal solutions while under-constraining unseen regions, leading to the Janus problem, mirrored faces, redundant limbs, and structurally inconsistent backsides [2], [16], [17]. Texture-oriented systems such as TextMesh [18] may improve appearance, yet they can still inherit underlying structural ambiguity because the optimization signal remains dominated by 2D image priors."
    )
    add_body(
        doc,
        "From a geometric perspective, these failures matter because they are not just perceptual artifacts; they often correspond to measurable defects in topology and mesh regularity. For a closed orientable surface, the Euler characteristic and genus can be related as:"
    )
    add_equation(doc, "chi = V - E + F,    g = (2 - chi) / 2")
    add_body(
        doc,
        "When hallucinations create extra holes, disconnected parts, or duplicated structures, these invariants may change even if a single rendered view still appears plausible. This is why AGD emphasizes geometric grounding rather than relying only on image-level quality assessment."
    )
    add_body(
        doc,
        "Another challenge arises when multimodal models are introduced as detectors. Hallo3D [6] and vision-language backbones such as LLaVA [19] make it possible to describe visible anomalies in natural language, but multimodal models can also hallucinate, omit subtle defects, or overstate spurious ones [20]. When the output of a detector is converted into a negative prompt instead of a localized 3D constraint, the system creates a semantic bottleneck: the geometry is corrected indirectly through language, rather than directly through spatial evidence. This remains one of the clearest unsolved problems in the literature."
    )

    add_heading(doc, "2.3 Comparative Analysis of Existing Methods", level=1)
    add_body(
        doc,
        "The reviewed studies reveal a progression from representation-focused systems to correction-aware systems, but they also reveal a repeated pattern: methods that generate attractive views do not necessarily offer reliable structural correction. Table 2.1 compares representative models in terms of representation, strengths, limitations, and their specific relevance to AGD."
    )
    add_comparison_table(doc)
    add_body(
        doc,
        "Table 2.1 shows that no single prior approach closes the full loop that AGD requires. Early explicit methods expose geometry clearly but lack semantics [7], [8]; SDS-based pipelines offer strong semantics but weak topology control [2]-[5], [12]; and multimodal correction methods improve diagnosis yet still depend on coarse textual feedback [6], [19], [20]."
    )

    add_heading(doc, "2.4 Strengths and Limitations of Current Approaches", level=1)
    add_body(
        doc,
        "Explicit and semi-explicit methods remain valuable because they keep geometric primitives interpretable. Point-based and voxel-based systems expose occupancy and shape errors more directly than latent image priors, which is useful for analysis and downstream editing [7], [8], [10]. Their weakness is that they often sacrifice surface continuity, fine detail, or text-level semantic richness, making them poor standalone solutions for high-fidelity text-to-3D synthesis."
    )
    add_body(
        doc,
        "Implicit-field and SDS-based methods constitute the strongest current family for open-ended 3D generation. NeRF [1] provides a continuous representation, while DreamFusion [2], Latent-NeRF [12], Magic3D [3], Fantasia3D [4], and ProlificDreamer [5] progressively improve texture quality, resolution, and diversity. Their shared limitation is structural: the optimization remains anchored to 2D consistency rather than native 3D reasoning. As a result, these models can generate visually convincing outputs that still contain duplicated geometry, disconnected components, or unstable topology."
    )
    add_body(
        doc,
        "Regularization and mitigation methods narrow specific failure modes but rarely solve the whole problem. ESD [16] and Perp-Neg [17] reduce mode collapse by modifying the guidance signal, while symmetry-aware or mesh-native approaches such as SYM3D [15] and MeshGPT [13] inject stronger structural priors into the representation itself. These are important advances because they move the field closer to geometric reasoning, yet they still do not provide a general mechanism for detecting errors across rendered views and correcting them inside the 3D object that caused them."
    )
    add_body(
        doc,
        "Multimodal correction systems contribute another important strength: they make failure analysis more interpretable. Hallo3D [6] demonstrates that a detector can express structurally meaningful complaints, and LLaVA-style models [19] expand the feasibility of local vision-language inspection. The weakness is that multimodal reasoning is itself imperfect [20]. If the detector is wrong, vague, or only partially grounded, the correction stage may over-correct, create fresh artifacts, or miss the most important region altogether. For AGD, this implies that semantic detection must be paired with a geometric critic and a spatial grounding module."
    )
    add_body(
        doc,
        "A related limitation lies in the nature of common mesh regularizers. Many pipelines rely on local smoothness priors of the form:"
    )
    add_equation(doc, "L_geo = sum_i || v_i - mean(N(i)) ||^2")
    add_body(
        doc,
        "Such terms are useful for suppressing noise and isolated spikes, but they do not explicitly encode topological validity, inter-view consistency, or anomaly localization. A smoother mesh is not necessarily a correct mesh. This distinction is central to why AGD frames refinement as grounded geometric correction rather than generic denoising."
    )

    add_heading(doc, "2.5 Summary of Issues", level=1)
    add_body(
        doc,
        "Across the literature, four issues recur. First, the supervision signal is frequently 2D and therefore under-constrains hidden geometry [2], [12], [18]. Second, mitigation methods often reshape the score space without verifying whether the underlying 3D structure has improved [16], [17]. Third, strong-looking outputs can still hide topological defects that are not obvious from a small set of rendered views [1], [5], [14]. Fourth, multimodal detectors improve explainability but introduce their own hallucination and grounding risks [6], [19], [20]."
    )
    add_body(
        doc,
        "These issues explain why better generation alone has not solved structural hallucinations. What is missing is a way to connect observed failures back to the specific geometric regions that produced them, and then refine those regions with an objective that respects both appearance and structure."
    )

    add_heading(doc, "2.6 Discussion on Research Gaps", level=1)
    add_body(
        doc,
        "The core research gap is not the absence of high-capacity generators; it is the absence of a closed-loop correction mechanism that translates observed hallucinations into spatially precise 3D constraints. Prior work either improves the generator, modifies the score signal, or adds a language-based detector [2]-[6], [16], [17]. Very few approaches jointly combine multi-view evidence, localized geometric criticism, and direct refinement of the 3D representation that caused the visible anomaly."
    )
    add_body(
        doc,
        "AGD is motivated by the need to close that loop. In the proposed framing of this study, the refinement objective augments the base generation loss with a geometry-aware energy:"
    )
    add_equation(doc, "L_total = L_SDS + E(x),    E(x) = lambda L_geo + (1 - lambda) D(x)")
    add_body(
        doc,
        "Here, the geometric term is not intended as a generic smoothing penalty alone; it is intended to be activated by grounded evidence from rendered views, auxiliary critics, and structural metrics. This differs from prior mitigation strategies because it does not ask text to stand in for geometry. Instead, it seeks to map view-level anomalies back into local 3D regions and to correct them through direct optimization."
    )

    add_heading(doc, "2.6.1 Definition of the research gap/problem", level=2)
    add_body(
        doc,
        "The research problem addressed by this study can therefore be defined as follows: current text-to-3D systems lack a reliable mechanism for detecting structural hallucinations across multiple views, grounding those detections to the responsible 3D regions, and refining the geometry through topology-aware spatial constraints rather than imprecise negative prompts. A successful AGD framework must satisfy all three requirements simultaneously: detection, grounding, and correction."
    )

    add_heading(doc, "2.7 Summary", level=1)
    add_body(
        doc,
        "This chapter reviewed the development of text-to-3D generation from explicit geometry and implicit fields to diffusion-guided synthesis, mesh-native modeling, and multimodal correction. The review showed that recent methods have improved resolution, semantics, and diversity, but they still struggle to guarantee structural consistency, especially under hidden viewpoints and topology-sensitive conditions."
    )
    add_body(
        doc,
        "The comparative analysis and gap discussion make the motivation for AGD explicit: the field now needs a refinement layer that can transform multi-view evidence and auxiliary criticism into direct geometric correction. The next chapter therefore presents the proposed AGD methodology, including the detection pathway, geometric grounding mechanism, and refinement objective used to mitigate hallucinations in text-to-3D generation."
    )

    add_heading(doc, "References", level=1)
    references = [
        "[1] B. Mildenhall, P. P. Srinivasan, M. Tancik, J. T. Barron, R. Ramamoorthi, and R. Ng, \"NeRF: Representing scenes as neural radiance fields for view synthesis,\" Commun. ACM, vol. 65, no. 1, pp. 99-106, Jan. 2022, doi: 10.1145/3503250.",
        "[2] B. Poole, A. Jain, J. T. Barron, and B. Mildenhall, \"DreamFusion: Text-to-3D using 2D Diffusion,\" arXiv:2209.14988, 2022, doi: 10.48550/arXiv.2209.14988.",
        "[3] C.-H. Lin et al., \"Magic3D: High-Resolution Text-to-3D Content Creation,\" arXiv:2211.10440, 2023, doi: 10.48550/arXiv.2211.10440.",
        "[4] R. Chen, Y. Chen, N. Jiao, and K. Jia, \"Fantasia3D: Disentangling Geometry and Appearance for High-quality Text-to-3D Content Creation,\" arXiv:2303.13873, 2023, doi: 10.48550/arXiv.2303.13873.",
        "[5] Z. Wang et al., \"ProlificDreamer: High-Fidelity and Diverse Text-to-3D Generation with Variational Score Distillation,\" arXiv:2305.16213, 2023, doi: 10.48550/arXiv.2305.16213.",
        "[6] H. Wang, J. Cao, J. Liu, X. Zhou, H. Huang, and R. He, \"Hallo3D: Multi-Modal Hallucination Detection and Mitigation for Consistent 3D Content Generation,\" in Proc. NeurIPS, 2024.",
        "[7] H. Fan, H. Su, and L. Guibas, \"A Point Set Generation Network for 3D Object Reconstruction from a Single Image,\" in Proc. CVPR, 2017.",
        "[8] K. Schwarz et al., \"VoxGRAF: Fast 3D Aware Image Synthesis with Sparse Voxel Grids,\" in Proc. NeurIPS, 2022.",
        "[9] A. Jain et al., \"Zero-Shot Text-Guided Object Generation with Dream Fields,\" arXiv:2112.01455, 2022.",
        "[10] A. Nichol et al., \"Point-E: A System for Generating 3D Point Clouds from Complex Prompts,\" arXiv:2212.08751, 2022.",
        "[11] H. Jun and A. Nichol, \"Shap-E: Generating Conditional 3D Implicit Functions,\" arXiv:2305.02463, 2023.",
        "[12] G. Metzer, E. Richardson, O. Patashnik, R. Giryes, and D. Cohen-Or, \"Latent-NeRF for Shape-Guided Generation of 3D Shapes and Textures,\" arXiv:2211.07600, 2022, doi: 10.48550/arXiv.2211.07600.",
        "[13] Y. Siddiqui et al., \"MeshGPT: Generating Triangle Meshes with Decoder-Only Transformers,\" in Proc. CVPR, 2024.",
        "[14] T. Wu et al., \"Recent advances in 3D Gaussian splatting,\" Comput. Vis. Media, vol. 10, no. 4, pp. 613-642, Aug. 2024, doi: 10.1007/s41095-024-0436-y.",
        "[15] J. Yang et al., \"SYM3D: Learning Symmetric Triplanes for Better 3D Awareness of GANs,\" arXiv:2406.06432, 2024.",
        "[16] P. Wang et al., \"Taming Mode Collapse in Score Distillation for Text-to-3D Generation,\" arXiv:2401.00909, 2024, doi: 10.48550/arXiv.2401.00909.",
        "[17] M. Armandpour, A. Sadeghian, H. Zheng, A. Sadeghian, and M. Zhou, \"Re-imagine the Negative Prompt Algorithm: Transform 2D Diffusion into 3D, alleviate Janus problem and Beyond,\" arXiv:2304.04968, 2023, doi: 10.48550/arXiv.2304.04968.",
        "[18] C. Tsalicoglou, F. Manhardt, A. Tonioni, M. Niemeyer, and F. Tombari, \"TextMesh: Generation of Realistic 3D Meshes From Text Prompts,\" arXiv:2304.12439, 2023, doi: 10.48550/arXiv.2304.12439.",
        "[19] H. Liu, C. Li, Q. Wu, and Y. J. Lee, \"Visual Instruction Tuning,\" arXiv:2304.08485, 2023.",
        "[20] Z. Bai et al., \"Hallucination of Multimodal Large Language Models: A Survey,\" arXiv:2404.18930, 2025, doi: 10.48550/arXiv.2404.18930.",
    ]
    for entry in references:
        add_reference_entry(doc, entry)

    output_path = Path.cwd() / OUTPUT_NAME
    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    path = build_document()
    print(f"Saved {path}")
