"""Generate AGD Research Word Document."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────────
def heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri+1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()

# ── Cover Page ───────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("\n\n")

title = doc.add_heading("Adversarial Geometric Distillation (AGD)\nResearch Documentation", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run("N D G S Dissasekara | 215519C\n").bold = True
info.add_run("BSc (Hons) in Artificial Intelligence\n")
info.add_run("Department of Computational Mathematics\n")
info.add_run("Faculty of Information Technology, University of Moratuwa\n")
info.add_run(f"Date: {datetime.date.today().strftime('%B %Y')}\n")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 1: APPROACH
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "Chapter 1: Research Approach", 1)

body(doc,
    "This chapter presents the complete research approach for Adversarial Geometric Distillation (AGD) — "
    "a closed-loop geometric grounding framework designed to mitigate hallucinations in text-to-3D "
    "model generation. It covers the hypothesis, inputs, processing pipeline, outputs, key features, "
    "and the technologies employed.")

# ── 1.1 Hypothesis ───────────────────────────────────────────────────────────
heading(doc, "1.1 Hypothesis", 2)
body(doc,
    "Current text-to-3D generation frameworks (e.g., DreamFusion, Magic3D, Fantasia3D) suffer from "
    "geometric hallucinations — most prominently the 'Janus problem' — because they rely on 2D "
    "diffusion priors that lack inherent 3D spatial awareness. Existing mitigation methods such as "
    "Hallo3D use Large Multimodal Models (LMMs) to detect structural errors, but their textual "
    "feedback creates a 'semantic bottleneck': negative text prompts cannot carry the spatial "
    "precision needed to correct 3D geometry, often causing over-correction or new artifacts.")
body(doc,
    "Core Hypothesis: By replacing the text-based correction loop with a direct, "
    "spatially-aware geometric feedback mechanism — an Adversarial Geometric Critic grounded in "
    "topological and spectral 3D metrics — it is possible to eliminate both primary 2D-prior "
    "hallucinations and secondary LMM reasoning errors, yielding structurally sound, "
    "manifold-consistent 3D meshes without re-prompting.")

# ── 1.2 Input ────────────────────────────────────────────────────────────────
heading(doc, "1.2 Inputs", 2)
body(doc, "The AGD pipeline accepts the following inputs:")
add_table(doc,
    ["Input Type", "Format", "Description"],
    [
        ["Text Prompt", "String", "Natural language description of desired 3D object (e.g., 'a dragon')"],
        ["3D Mesh (prototype)", ".obj / .ply / .stl", "Pre-generated mesh files used for validation in the prototype"],
        ["Camera Parameters", "Azimuth & Elevation (°)", "Defines multi-view rendering angles (6 canonical or up to 100+ via Fibonacci sphere)"],
        ["Hyperparameters", "CLI flags", "lr, iterations, lambda_geo, lambda_distill, num_views, view_method, n_sample_pts"],
    ])

# ── 1.3 Process ──────────────────────────────────────────────────────────────
heading(doc, "1.3 Process / Pipeline Architecture", 2)
body(doc, "The AGD pipeline consists of six sequential stages:")

stages = [
    ("Stage 1 — Generative Synthesis", "SDS-based generation of an initial 3D representation "
     "(SDF or mesh) from a text/image input using 2D diffusion priors (DreamFusion/Threestudio). "
     "In the current prototype, this stage is simulated by loading existing mesh files."),
    ("Stage 2 — Multi-View Rendering", "The mesh is rendered from N mathematically distributed "
     "camera positions. The upgraded renderer uses Fibonacci-sphere sampling (golden-angle spiral: "
     "φ = arccos(1-2(i+0.5)/n), θ = 2πi/φ_gold) to achieve uniform angular coverage. "
     "Each view produces 5 image buffers: RGB, depth (perspective-correct 1/z), interpolated "
     "normal map, binary silhouette, and curvature edge map."),
    ("Stage 3 — LMM Hallucination Detection", "A local LLaVA vision-language model inspects each "
     "rendered view and returns a JSON severity score + notes (e.g., {\"severity\": 0.7, "
     "\"notes\": \"duplicated face detected\"}). View-consistency is also measured via "
     "feature-similarity across views."),
    ("Stage 4 — Adversarial Geometric Critic", "A heuristic critic (with optional GNN) computes "
     "per-vertex anomaly scores from degree z-scores and edge-length outliers. Spectral metrics "
     "(Laplacian eigenvalue variance, Fiedler value) and topological metrics (Euler characteristic, "
     "genus, connected components) are also computed. The critic identifies which spatial regions "
     "of the mesh contain hallucinated geometry."),
    ("Stage 5 — Geometric Grounding", "View-level severity scores and critic scores are combined "
     "into a spatially-aware vertex weight map w_i. Each view direction is mapped to vertices "
     "lying near the corresponding extent of the mesh and facing the view normal. This replaces "
     "the 'semantic bottleneck' with a precise 3D spatial constraint."),
    ("Stage 6 — Grounded Refinement Loop", "Iterative vertex optimization using a combined energy:\n"
     "    E(x) = λ·L_geo + (1-λ)·D\n"
     "where L_geo = uniform Laplacian smoothing (Σ||v_i - mean(N(i))||) and "
     "D = distillation anchor term (Σ||v_i - v_i^(0)||). Updates are applied only in grounded "
     "high-weight regions: v ← v - η·(λ∇L_geo + (1-λ)∇D)⊙w"),
]
for title_s, desc in stages:
    p = doc.add_paragraph()
    p.add_run(title_s + ": ").bold = True
    p.add_run(desc)
    p.paragraph_format.space_after = Pt(6)

# ── 1.4 Output ───────────────────────────────────────────────────────────────
heading(doc, "1.4 Outputs", 2)
add_table(doc,
    ["Output", "Format", "Description"],
    [
        ["Refined 3D Mesh", "*_agd.obj/.ply/.stl", "Structurally corrected mesh with reduced hallucinations"],
        ["Multi-view RGB images", "PNG (per view)", "Phong-shaded renders from all N camera positions"],
        ["Depth maps", "PNG (grayscale)", "Perspective-correct depth buffer per view"],
        ["Normal maps", "PNG (RGB)", "Per-pixel interpolated vertex normals per view"],
        ["Silhouette masks", "PNG (binary)", "Object boundary mask per view"],
        ["Edge/curvature maps", "PNG (binary)", "Sharp discontinuity detection per view"],
        ["Geometry Error Report", "Console + dict", "Chamfer, Hausdorff, NCS, Angular Error, Roughness RMS, SSIM, Edge-IoU"],
        ["Hallucination Score", "Float [0,1]", "Before/after composite score from discriminator"],
        ["View Metadata JSON", "JSON", "az/el angles for all rendered views"],
    ])

# ── 1.5 Features ─────────────────────────────────────────────────────────────
heading(doc, "1.5 Key Features", 2)

features = [
    ("Closed-Loop Geometric Grounding",
     "Unlike Hallo3D's text prompts, AGD feeds LMM detections back as spatial energy functions "
     "directly into the 3D optimization — no semantic bottleneck."),
    ("Dense Multi-View Scanning (6→100+ views)",
     "Fibonacci-sphere camera sampling distributes up to 100 views uniformly on the sphere "
     "(golden-angle: no pole clustering). More views = higher chance of detecting hidden hallucinations."),
    ("Advanced Rendering Pipeline",
     "Phong shading (ambient+diffuse+specular), perspective-correct 1/z depth interpolation, "
     "per-pixel barycentric vertex-normal interpolation, and curvature edge detection."),
    ("Multi-Metric Error Quantification",
     "10 geometry metrics: Chamfer Distance, Hausdorff (HD95), Normal Consistency, Angular Normal "
     "Error (degrees), Surface Roughness RMS, Curvature Wasserstein-1, Silhouette IoU, Depth RMSE, "
     "SSIM, Edge-IoU. Multi-view aggregate (min/max/std/p5/p95 across all N views)."),
    ("Dual Critic Architecture",
     "Heuristic critic (degree + edge-length z-scores) works out-of-the-box; optional GNN critic "
     "(Graph Convolutional Network) can be trained with pseudo-labels from the heuristic."),
    ("Topology-Aware Scoring",
     "Discriminator computes Euler characteristic, genus, Betti numbers, Laplacian variance, "
     "and Fiedler value — metrics that directly encode structural hallucination severity."),
    ("Single-File or Batch Processing",
     "Run on one mesh with --input-file or an entire folder with --input-dir."),
]
for feat, desc in features:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(feat + ": ").bold = True
    p.add_run(desc)

# ── 1.6 Technologies ─────────────────────────────────────────────────────────
heading(doc, "1.6 Technologies & Tools", 2)
add_table(doc,
    ["Category", "Technology", "Role"],
    [
        ["3D Geometry", "Trimesh", "Mesh loading, topology, surface sampling, face/vertex normals"],
        ["3D Rendering", "NumPy (custom rasteriser)", "Software rasteriser — no GPU/display server needed"],
        ["Spectral Analysis", "SciPy (sparse, cKDTree)", "Laplacian eigenvalue estimation, KD-tree nearest-neighbour"],
        ["Graph ML", "PyTorch + PyTorch Geometric", "GNN critic (GCNConv layers) for vertex anomaly scoring"],
        ["VLM Detection", "LLaVA v1.5/1.6 (local)", "Per-view hallucination severity detection"],
        ["Image Processing", "Pillow, scikit-image", "PNG I/O, SSIM computation"],
        ["Graph Analysis", "NetworkX", "Connected component counting for topology metrics"],
        ["3D Generation (planned)", "Threestudio + PyTorch", "SDS-based text-to-3D synthesis backbone"],
        ["Datasets (planned)", "Objaverse, GSO", "Benchmark evaluation of geometric quality"],
        ["Hardware Target", "NVIDIA RTX 3090/4090 (24GB+)", "LMM inference + differentiable 3D optimization"],
    ])

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER 2: ACHIEVEMENTS, LIMITATIONS & TESTS
# ════════════════════════════════════════════════════════════════════════════
heading(doc, "Chapter 2: Achievements, Limitations & Tested Results", 1)

# ── 2.1 What Has Been Achieved ───────────────────────────────────────────────
heading(doc, "2.1 What Has Been Achieved", 2)

body(doc, "The following components of the proposed AGD framework have been fully implemented and validated in the prototype:")

achievements = [
    ("✅ Closed-Loop Refinement Pipeline",
     "The full 6-stage pipeline (load → render → detect → critic → ground → optimize → evaluate) "
     "runs end-to-end on real mesh files. Successfully processed 13+ mesh assets."),
    ("✅ Multi-View Renderer (v2) — 6 to 100+ Views",
     "Upgraded from 6 fixed canonical views to arbitrary N views via Fibonacci-sphere or "
     "azimuth-elevation grid. Phong shading, perspective-correct depth, interpolated normals, "
     "and edge map buffers — 5 image types per view. Verified: 100-view Fibonacci = 100 views, "
     "500 PNG files per mesh."),
    ("✅ Geometric Error Metrics — 10 Metrics",
     "All 10 metrics implemented and tested: Chamfer Distance, Hausdorff (+ HD95), Normal "
     "Consistency Score, Angular Normal Error (°), Local Surface Roughness RMS, Curvature "
     "Wasserstein-1, Silhouette IoU, Depth RMSE, SSIM, Edge-IoU. Multi-view aggregate "
     "statistics (min/max/std/p5/p95) computed across all N views."),
    ("✅ Adversarial Geometric Critic (Heuristic)",
     "Per-vertex anomaly scoring using degree z-scores and edge-length z-scores. Verified "
     "on all dataset meshes. GNN stub implemented with GCNConv and training script (gnn_train.py)."),
    ("✅ LMM Hallucination Detector (LLaVA)",
     "Local LLaVA v1.5 integration — per-view severity JSON parsing, global bias computation, "
     "and view-score propagation into vertex weights."),
    ("✅ Geometric Grounding Module",
     "View-level severity → spatially-aware vertex weight map. Direction-based vertex selection "
     "(quantile-based extent + normal facing). Combines critic scores and LMM bias into final w_i."),
    ("✅ Hallucination Discriminator",
     "Topology metrics: Euler characteristic, genus, Betti numbers, connected components. "
     "Spectral metrics: Laplacian eigenvalue variance, Fiedler value (algebraic connectivity). "
     "Composite hallucination score [0,1] with 4-weight penalty formula."),
    ("✅ Single-File Processing",
     "Added --input-file flag for processing one mesh at a time alongside --input-dir batch mode."),
    ("✅ Dataset Analysis — 13 Meshes",
     "Full pipeline tested on: bunny.ply, dragon.ply, janus.obj, Mr_Bean.obj, "
     "Meshy_AI_Infernal_Ironclad.obj, hand.ply, horse.ply, elepham.obj, venusm.obj, "
     "mba1.obj, mba2.obj, ateneam.obj, Rigged_Hand.obj."),
]
for feat, desc in achievements:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(feat + "\n").bold = True
    p.add_run(desc)
    p.paragraph_format.space_after = Pt(4)

# ── 2.2 Test Results ─────────────────────────────────────────────────────────
heading(doc, "2.2 Tested Results with Examples", 2)

body(doc, "All 13 meshes in the dataset were processed. Key results from the pipeline run "
     "(python agd_pipeline.py --input-dir 3d_samples --render-views) are shown below:")

add_table(doc,
    ["Mesh", "Score", "Variance", "Fiedler", "Interpretation"],
    [
        ["ateneam.obj", "0.878", "2.892", "≈0", "HIGH — severe connectivity irregularity, spiky protrusions"],
        ["bunny.ply", "0.488", "0.676", "≈0", "MODERATE — open boundary at base causes disconnection"],
        ["dragon.ply", "0.693", "0.039", "≈0", "MOD-HIGH — smooth mesh but multiple components"],
        ["elepham.obj", "0.650", "1.057", "≈0", "MOD-HIGH — medium variance; disconnected parts"],
        ["hand.ply", "0.610", "0.061", "≈0", "MODERATE — clean geometry but sub-mesh disconnection"],
        ["horse.ply", "0.400", "0.519", "≈0", "LOW-MODERATE — relatively clean mesh"],
        ["janus.obj", "0.672", "1.162", "≈0", "MOD-HIGH — Janus artifact, duplicated face geometry"],
        ["mba1.obj", "1.000", "0.146", "≈0", "MAXIMUM — extreme topological complexity"],
        ["mba2.obj", "1.000", "0.189", "≈0", "MAXIMUM — extreme topological complexity"],
        ["Meshy_AI_Infernal_Ironclad.obj", "1.000", "0.053", "≈0", "MAXIMUM — AI mesh, many separate shell components"],
        ["Mr_Bean.obj", "1.000", "2.538", "≈0", "MAXIMUM + HIGH VAR — AI hallucination, extreme score"],
        ["Rigged_Hand.obj", "0.685", "7.333", "≈0", "HIGHEST VARIANCE — bone/skin overlap, extreme NN variation"],
        ["venusm.obj", "1.000", "1.063", "≈0", "MAXIMUM — high genus or many components"],
    ])

body(doc, "Multi-View Rendering Test (36-view grid):")
bullet(doc, "Command: python agd_pipeline.py --render-views --num-views 36 --view-method grid")
bullet(doc, "Result: 32 views rendered → 160 PNG image files generated per mesh")
bullet(doc, "Buffers per view: rgb, depth, normal_map, silhouette, edge_map")
bullet(doc, "Geometry metrics computed: Chamfer, Hausdorff, NCS, Angular error, Roughness, SSIM")

doc.add_paragraph()
body(doc, "100-View Fibonacci Test:")
bullet(doc, "Command: python agd_pipeline.py --input-file janus.obj --render-views --num-views 100 --view-method fibonacci")
bullet(doc, "Camera schedule verified: 100 views, uniform golden-angle spiral distribution")
bullet(doc, "First view: (az=0°, el=81.9°) — near north pole")
bullet(doc, "Last view: (az=66.7°, el=-81.9°) — near south pole, mirrored")

# ── 2.3 Limitations ──────────────────────────────────────────────────────────
heading(doc, "2.3 Current Limitations", 2)

limitations = [
    ("❌ No SDS-Based Text-to-3D Generation",
     "The pipeline currently requires pre-existing mesh files. The full SDS-based synthesis "
     "(text → 3D) using Threestudio/DreamFusion is not yet integrated. This means the pipeline "
     "acts as a post-processing refinement layer rather than an end-to-end generator."),
    ("❌ Topological Metrics Not Used as Differentiable Losses",
     "Genus, Euler characteristic, and Betti numbers are computed for scoring but are NOT used "
     "in the actual optimization loss. Only Laplacian smoothing and an anchor term drive refinement. "
     "True topology-editing operations (hole filling, component merging, remeshing) are absent."),
    ("❌ Scores Unchanged Before→After",
     "All 13 meshes showed identical hallucination scores before and after optimization. "
     "Laplacian smoothing shifts vertex positions but not topology — spectral/topological metrics "
     "are computed on the same adjacency graph and do not change with positional updates."),
    ("❌ No Precise VLM Spatial Localization",
     "LLaVA returns a single severity score per view, not bounding boxes or pixel masks. "
     "The grounding module maps view-level scores to vertices by direction heuristics "
     "(quantile-based extent), which is coarse and can bias the wrong vertices."),
    ("❌ GNN Critic Uses Only Pseudo-Labels",
     "The GNN (GCNConv) critic is trained using pseudo-labels from the heuristic critic "
     "(degree/edge-length z-scores) rather than real annotated hallucination examples. "
     "This limits its ability to generalize beyond the heuristic."),
    ("❌ Vertex Limit (500,000)",
     "Meshes exceeding 500,000 vertices are skipped to prevent out-of-memory errors. "
     "3 meshes (Bearded_guy.ply: 1,024,355 vertices; blade.ply: 882,954; happy.ply: 543,524) "
     "were skipped in the full dataset run."),
    ("❌ CPU-Bound Rendering is Slow",
     "The pure-numpy software rasteriser processes faces sequentially in Python. "
     "Rendering 100 views of a high-poly mesh (69,451 faces, Mr_Bean.obj) took several minutes. "
     "No GPU-accelerated rendering (e.g., nvdiffrast) is used."),
    ("❌ No Benchmark Evaluation (Objaverse/GSO)",
     "The pipeline has not been evaluated on standard benchmarks (Objaverse, Google Scanned Objects). "
     "Quantitative comparison against baseline methods (Hallo3D, DreamFusion) is pending."),
]
for feat, desc in limitations:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(feat + "\n").bold = True
    p.add_run(desc)
    p.paragraph_format.space_after = Pt(4)

# ── 2.4 Suggested Solutions ───────────────────────────────────────────────────
heading(doc, "2.4 Suggested Solutions & Future Work", 2)

solutions = [
    ("1. Integrate SDS-Based Generation (Threestudio)",
     "Connect the pipeline to Threestudio/DreamFusion as the front-end generator. Feed the "
     "AGD refinement loop directly after the initial SDS generation step. This enables true "
     "end-to-end text-to-clean-3D processing."),
    ("2. Differentiable Topology Losses",
     "Replace Laplacian-only smoothing with topology-aware losses: self-intersection penalty "
     "(penalize faces whose projections overlap), manifoldness constraint (penalize non-manifold "
     "edges), and genus reduction via edge collapse/flip operations."),
    ("3. Projection-Based Grounding (Ray Casting)",
     "Instead of direction heuristics, cast rays from the camera through detected anomaly regions "
     "(2D pixel coordinates from VLM bounding boxes) into 3D space to identify exact vertices — "
     "giving pixel-accurate spatial grounding."),
    ("4. GPU-Accelerated Rendering (nvdiffrast)",
     "Replace the pure-numpy rasteriser with nvdiffrast or PyTorch3D for GPU-accelerated "
     "differentiable rendering. This would allow rendering gradients to flow back into vertex "
     "positions (differentiable silhouette/depth loss)."),
    ("5. Stronger GNN Training with Real Annotations",
     "Collect ground-truth hallucination vertex labels from Objaverse (by comparing generated "
     "meshes to reference scans). Train the GNN critic on these real annotations for "
     "geometry-aware, generalizable anomaly detection."),
    ("6. Increase Optimizer Strength",
     "Increase --iterations (50-200) and use adaptive learning rates (Adam-style). Add "
     "topology-editing: trimesh hole_filling, component merging, and Quadric Error Metrics "
     "decimation to actually reduce genus and component count."),
    ("7. Structured VLM Output (Bounding Boxes/Masks)",
     "Switch from LLaVA free-text to a structured detection model (Grounded-SAM or LLaVA-1.6 "
     "with region prompting) that returns bounding boxes or segmentation masks for detected "
     "hallucinations, enabling precise 2D→3D grounding."),
    ("8. Benchmark Evaluation",
     "Run on 100+ meshes from Objaverse and Google Scanned Objects. Compare Chamfer Distance, "
     "Hausdorff, and hallucination scores against DreamFusion, Magic3D, and Hallo3D baselines. "
     "Report FID-3D and coverage metrics."),
]
for title_s, desc in solutions:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(title_s + "\n").bold = True
    p.add_run(desc)
    p.paragraph_format.space_after = Pt(4)

# ── 2.5 Pipeline Coverage Summary ────────────────────────────────────────────
heading(doc, "2.5 Pipeline Coverage Summary", 2)

add_table(doc,
    ["Proposal Component", "Status", "Detail"],
    [
        ["SDS-Based 3D Generation", "❌ Not Implemented", "Requires Threestudio integration"],
        ["Multi-View Renderer", "✅ Implemented (v2)", "6→100 views, Phong, edge map, 5 buffers"],
        ["LMM Hallucination Detector", "✅ Implemented", "LLaVA local inference, JSON severity parsing"],
        ["View Consistency Detector", "✅ Implemented", "Feature-similarity across views"],
        ["Adversarial Geometric Critic", "⚠️ Partial", "Heuristic works; GNN is pseudo-label scaffold"],
        ["Geometric Grounding", "✅ Implemented (coarse)", "Direction-based vertex weight mapping"],
        ["Grounded Refinement Loop", "✅ Implemented", "Laplacian + anchor term, weighted update"],
        ["Topological Metrics", "✅ Computed", "Euler χ, genus, components, Fiedler, variance"],
        ["Geometric Error Report", "✅ Implemented (v2)", "10 metrics including SSIM, Angular Error, Roughness"],
        ["Benchmark Evaluation", "❌ Not Done", "Objaverse/GSO evaluation pending"],
    ])

doc.add_page_break()

# ── Reference ─────────────────────────────────────────────────────────────────
heading(doc, "References", 1)

refs = [
    "[1] B. Mildenhall et al., 'NeRF: Representing scenes as neural radiance fields for view synthesis,' Commun. ACM, 2022.",
    "[2] B. Poole et al., 'DreamFusion: Text-to-3D using 2D Diffusion,' arXiv:2209.14988, 2022.",
    "[3] C.-H. Lin et al., 'Magic3D: High-Resolution Text-to-3D Content Creation,' arXiv:2211.10440, 2023.",
    "[4] R. Chen et al., 'Fantasia3D: Disentangling Geometry and Appearance,' arXiv:2303.13873, 2023.",
    "[5] P. Wang et al., 'Taming Mode Collapse in Score Distillation,' arXiv:2401.00909, 2024.",
    "[6] M. Armandpour et al., 'Re-imagine the Negative Prompt Algorithm,' arXiv:2304.04968, 2023.",
    "[7] H. Wang et al., 'Hallo3D: Multi-Modal Hallucination Detection and Mitigation for Consistent 3D Content Generation.'",
    "[8] Z. Bai et al., 'Hallucination of Multimodal Large Language Models: A Survey,' arXiv:2404.18930, 2025.",
]
for ref in refs:
    bullet(doc, ref)

# ── Save ──────────────────────────────────────────────────────────────────────
out = "AGD_Research_Documentation.docx"
doc.save(out)
print(f"Saved: {out}")
